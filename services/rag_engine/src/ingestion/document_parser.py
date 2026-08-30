import json
from pathlib import Path
from typing import List, Dict, Any
from src.ingestion.clean import clean_text

class DocumentParser:
    """Multi-format parser for PDF, DOCX, TXT, CSV, and JSON documents."""

    @staticmethod
    def parse_file(file_path: str) -> List[Dict[str, Any]]:
        """Parse file based on extension and return structured list of page/section dicts."""
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".pdf":
            return DocumentParser._parse_pdf(path)
        elif ext == ".docx":
            return DocumentParser._parse_docx(path)
        elif ext == ".json":
            return DocumentParser._parse_json(path)
        elif ext in (".txt", ".csv"):
            return DocumentParser._parse_plaintext(path)
        else:
            raise ValueError(f"Unsupported document format: {ext}")

    @staticmethod
    def _parse_pdf(path: Path) -> List[Dict[str, Any]]:
        """Extract text and page metadata using pypdfium2."""
        sections = []
        try:
            import pypdfium2 as pdfium
            pdf = pdfium.PdfDocument(path)
            for i, page in enumerate(pdf):
                text_page = page.get_textpage()
                extracted = text_page.get_text_range()
                cleaned = clean_text(extracted)
                if cleaned:
                    sections.append({
                        "page_number": i + 1,
                        "section_title": f"Page {i + 1}",
                        "content_text": cleaned,
                    })
        except Exception as e:
            # Fallback plain text read if PDF extraction fails
            sections.append({
                "page_number": 1,
                "section_title": "Document Root",
                "content_text": clean_text(path.read_text(errors="ignore")),
            })
        return sections

    @staticmethod
    def _parse_docx(path: Path) -> List[Dict[str, Any]]:
        """Extract headings and paragraph blocks from Word DOCX files."""
        sections = []
        try:
            import docx
            doc = docx.Document(path)
            current_title = "Document Root"
            current_text = []
            page_counter = 1

            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    if current_text:
                        sections.append({
                            "page_number": page_counter,
                            "section_title": current_title,
                            "content_text": clean_text("\n".join(current_text)),
                        })
                        current_text = []
                    current_title = para.text.strip() or "Untitled Section"
                else:
                    if para.text.strip():
                        current_text.append(para.text)

            if current_text:
                sections.append({
                    "page_number": page_counter,
                    "section_title": current_title,
                    "content_text": clean_text("\n".join(current_text)),
                })
        except Exception as e:
            sections.append({
                "page_number": 1,
                "section_title": "Document Root",
                "content_text": clean_text(path.read_text(errors="ignore")),
            })
        return sections

    @staticmethod
    def _parse_json(path: Path) -> List[Dict[str, Any]]:
        """Parse structured JSON technical specifications."""
        content = path.read_text(encoding="utf-8")
        data = json.loads(content)
        return [{
            "page_number": 1,
            "section_title": "JSON Specification Root",
            "content_text": clean_text(json.dumps(data, indent=2)),
        }]

    @staticmethod
    def _parse_plaintext(path: Path) -> List[Dict[str, Any]]:
        """Parse TXT or CSV plaintext files."""
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [{
            "page_number": 1,
            "section_title": "Document Content",
            "content_text": clean_text(text),
        }]
